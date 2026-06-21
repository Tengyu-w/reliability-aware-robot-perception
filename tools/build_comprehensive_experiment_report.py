import json
import os
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "comprehensive_experiment_report.docx"


def read_json(rel):
    path = ROOT / rel
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def read_text(rel, fallback=""):
    path = ROOT / rel
    if not path.exists():
        return fallback
    return path.read_text(encoding="utf-8", errors="replace")


def all_files(pattern):
    return sorted(ROOT.glob(pattern), key=lambda p: str(p).lower())


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Title"].font.name = "Calibri"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor.from_string("0B2545")


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "Reliability-Aware Sequential and 3D Robot Perception | Comprehensive Experiment Report"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("666666")


def add_title(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Reliability-Aware Sequential and 3D Robot Perception\n").bold = True
    r = p.add_run("综合实验数据、实验细节与博士申请证据报告")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"项目路径: {ROOT}\n").italic = True
    p.add_run(f"生成日期: {date.today().isoformat()}\n").italic = True
    p.add_run("用途: GitHub 展示、导师沟通、博士申请研究经历整理").italic = True


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        set_cell_shading(hdr[i], "F4F6F9")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_caption(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")
    return p


def add_image(doc, rel_path, caption, width=5.9):
    path = ROOT / rel_path
    if not path.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)
    return True


def fmt(v, digits=3):
    if isinstance(v, bool):
        return "true" if v else "false"
    if isinstance(v, float):
        return f"{v:.{digits}f}"
    return v


def metrics_overview_rows():
    items = [
        ("Synthetic 3D reliability", "outputs/robot_3d_demo/robot_3d_metrics.json", "embedding_risk_roc_auc_for_corruption_detection"),
        ("Synthetic 3D multi-seed", "outputs/robot_3d_multiseed/robot_3d_multiseed_metrics.json", "embedding_risk_roc_auc_mean"),
        ("Smoke real-depth corruption", "outputs/real_depth_corruption_smoke/real_depth_corruption_metrics.json", "embedding_risk_roc_auc_for_corruption_detection"),
        ("Prepared-depth pipeline smoke", "outputs/robot_3d_pipeline_smoke/corruption_benchmark/real_depth_corruption_metrics.json", "embedding_risk_roc_auc_for_corruption_detection"),
        ("TUM RGB-D source-paired corruption", "outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_metrics.json", "embedding_risk_roc_auc_for_corruption_detection"),
        ("TUM scene-conditioned baseline", "outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_metrics.json", "scene_conditioned_roc_auc_for_corruption_detection"),
        ("TUM temporal excess scoring", "outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_metrics.json", "temporal_excess_roc_auc_for_corruption_detection"),
        ("TUM temporal local-distance baseline", "outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_metrics.json", "temporal_local_roc_auc_for_corruption_detection"),
        ("PCA descriptor rotation sensitivity", "outputs/tum_rgbd_freiburg1_desk_pca_descriptor/tum_pca_depth_descriptor_metrics.json", "spearman_rotation_vs_embedding_shift"),
        ("Calibration ROC-AUC", "outputs/tum_rgbd_calibration/calibration_metrics.json", "roc_auc"),
        ("Trajectory residual ROC-AUC", "outputs/trajectory_residual_demo/trajectory_residual_metrics.json", "roc_auc"),
    ]
    rows = []
    for name, rel, key in items:
        data = read_json(rel)
        rows.append([name, key, fmt(data.get(key, "n/a")), rel])
    return rows


def add_csv_preview(doc, rel_path, title, max_rows=12):
    path = ROOT / rel_path
    if not path.exists():
        return
    doc.add_heading(title, level=3)
    add_para(doc, f"CSV 文件: {rel_path}")
    df = pd.read_csv(path)
    add_para(doc, f"行数: {len(df)}; 列数: {len(df.columns)}; 列名: {', '.join(map(str, df.columns[:12]))}" + (" ..." if len(df.columns) > 12 else ""))
    shown = df.head(max_rows).copy()
    if shown.shape[1] > 6:
        shown = shown.iloc[:, :6]
    rows = shown.fillna("").astype(str).values.tolist()
    add_table(doc, list(shown.columns), rows, widths=[int(9360 / max(1, shown.shape[1]))] * shown.shape[1])


def add_markdown_report_appendix(doc, rel_path):
    text = read_text(rel_path)
    if not text:
        return
    doc.add_heading(rel_path, level=3)
    lines = [ln.strip() for ln in text.splitlines()]
    for ln in lines:
        if not ln:
            continue
        if ln.startswith("# "):
            add_para(doc, ln[2:], style=None)
        elif ln.startswith("## "):
            p = add_para(doc, ln[3:])
            p.runs[0].bold = True
        elif ln.startswith("- "):
            doc.add_paragraph(ln[2:], style="List Bullet")
        else:
            add_para(doc, ln)


def build():
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_title(doc)

    doc.add_page_break()
    doc.add_heading("阅读说明", level=1)
    add_para(doc, "这份综合报告把当前项目已经产生的实验内容、实验数据、实验流程、实验规律、图片证据和局限性整合到一个 Word 文件中。报告以博士申请和导师沟通为目标，强调研究问题、证据链、可复现实验流程和下一步研究方向。")
    add_para(doc, "需要注意：部分 CSV 包含上千行样本级数据。为了让 Word 文件可打开、可阅读，正文放入关键统计、代表性预览和全部图片；完整 CSV、JSON、Markdown 子报告的路径在附录中列出，原文件仍保留在项目目录。")

    doc.add_heading("目录", level=1)
    toc = [
        "1. 项目总体定位",
        "2. 代码和数据结构",
        "3. 实验总览与关键结果表",
        "4. 视频序列模型与 embedding 诊断",
        "5. 合成 3D 深度感知可靠性实验",
        "6. 真实深度图准备与 smoke pipeline",
        "7. TUM RGB-D 公开数据实验",
        "8. 时序可靠性、姿态分析和 PCA 描述子",
        "9. 校准、coverage-risk 与运行时监控",
        "10. 轨迹残差与执行可靠性实验",
        "11. 实验规律、博士申请价值和下一步计划",
        "12. 附录: 指标、CSV 预览、图片索引和子报告",
    ]
    add_numbered(doc, toc)

    doc.add_heading("1. 项目总体定位", level=1)
    add_para(doc, "本项目当前的核心不是单纯追求一个分类器的最高准确率，而是围绕“可靠性”组织实验：当一个学习系统面对视频序列、深度图、点云几何或机器人轨迹时，如何判断当前输入、表示或执行结果是否已经偏离正常模式。")
    add_para(doc, "项目起点是 ResNet18 + LSTM 视频动作识别 baseline。随后项目被扩展为一个 reliability-aware perception and action monitoring prototype，覆盖 sequential perception、RGB-D robot perception、runtime assurance、calibration、selective prediction 和 trajectory residual monitoring。")
    add_bullets(doc, [
        "研究问题: embedding-space evidence、校准分析和残差监控能否识别学习系统的不可靠状态。",
        "应用方向: trustworthy ML、robot perception、embodied AI、safe autonomy、runtime assurance、safe RL、medical/surgical robotics。",
        "当前状态: 研究原型和 evidence pack；有公开数据 TUM RGB-D 实验，有合成和 smoke 实验，也有明确局限性说明。",
    ])

    doc.add_heading("2. 代码和数据结构", level=1)
    add_table(doc, ["目录/文件", "作用"], [
        ["modules/model.py", "定义 ResNet18 + LSTM 视频动作识别模型，并支持提取 embedding。"],
        ["modules/main.py", "训练入口: 加载 AVI 数据、训练、验证、预测、保存模型和指标。"],
        ["modules/embedding_analysis.py", "提取验证集 embedding、置信度、熵、margin、PCA 图和分类报告。"],
        ["modules/robot_3d_reliability.py", "深度图转点云、几何 embedding、深度腐蚀、风险评分、coverage-risk、TUM pose 分析相关工具。"],
        ["modules/runtime_monitor.py", "把连续风险分数转成 NORMAL/SUSPECT/RECOVER/HUMAN_REVIEW 状态。"],
        ["modules/calibration_analysis.py", "风险分数校准、ECE-style gap 和 coverage-risk 分析。"],
        ["modules/trajectory_residual_demo.py", "计划轨迹 vs 实际轨迹的残差可靠性实验。"],
        ["docs/", "项目概览、导师一页纸、申请证据包、数据卡模板等文档。"],
        ["outputs/", "实验结果、CSV、JSON、Markdown 报告和图片。"],
    ], widths=[2600, 6760])

    doc.add_heading("3. 实验总览与关键结果表", level=1)
    add_para(doc, "下表整合了当前项目中最重要的实验指标。ROC-AUC 越接近 1，表示风险分数越能把正常样本和异常/失败样本排序分开；Spearman 相关越高，表示 embedding shift 与真实相机运动或轨迹变化的单调关系越强。")
    add_table(doc, ["实验层", "指标字段", "结果", "来源文件"], metrics_overview_rows(), widths=[2300, 2500, 1100, 3460])
    add_para(doc, "关键读法: TUM source-paired 和 temporal excess 的 ROC-AUC 为 1.000，说明在受控深度腐蚀设置中，源图像配对或局部时间归一化可以非常清楚地识别异常；但 scene-conditioned baseline 约为 0.483，说明只用全局干净参考会被正常相机运动干扰。")

    doc.add_heading("4. 视频序列模型与 Embedding 诊断", level=1)
    add_para(doc, "视频部分使用 ResNet18 作为每一帧的 CNN 特征提取器，再用 LSTM 建模帧序列。模型最后输出动作类别，同时可以返回 LSTM 最后时间步的 hidden embedding。")
    add_bullets(doc, [
        "输入形式: batch, sequence, 3-channel image, height, width。",
        "CNN 输出维度: ResNet18 的 512 维视觉特征。",
        "LSTM hidden dimension: 默认 256。",
        "分类头: Dropout + Linear + ReLU + Linear。",
        "诊断指标: accuracy、classification report、confusion matrix、confidence、entropy、margin、embedding PCA。",
    ])
    add_para(doc, "该部分对博士申请的价值在于: 它证明项目并不是停留在静态图像分类，而是从 sequential perception 出发，具备处理时间序列模型和内部表示分析的基础。")
    if (ROOT / "modules" / "training_history.png").exists():
        add_image(doc, "modules/training_history.png", "图 1. 视频 baseline 训练历史曲线。")
    if (ROOT / "modules" / "test_predictions.csv").exists():
        add_csv_preview(doc, "modules/test_predictions.csv", "视频测试预测文件预览", max_rows=10)

    doc.add_heading("5. 合成 3D 深度感知可靠性实验", level=1)
    synth = read_json("outputs/robot_3d_demo/robot_3d_metrics.json")
    multi = read_json("outputs/robot_3d_multiseed/robot_3d_multiseed_metrics.json")
    add_para(doc, "合成 3D 实验的目的，是在不依赖真实机器人硬件的情况下，先验证一条完整的可靠性证据链: 生成深度场景、施加受控腐蚀、提取点云几何 embedding、用 clean reference 计算风险分数、评估高风险分数是否对应腐蚀样本。")
    add_table(doc, ["项目", "数值"], [
        ["合成 demo 样本数", synth.get("samples")],
        ["干净样本数", synth.get("clean_samples")],
        ["腐蚀样本数", synth.get("corrupted_samples")],
        ["单次 demo ROC-AUC", fmt(synth.get("embedding_risk_roc_auc_for_corruption_detection"))],
        ["单次 demo coverage-risk AUC", fmt(synth.get("coverage_risk_auc"))],
        ["multi-seed seeds", multi.get("seeds")],
        ["multi-seed ROC-AUC mean ± std", f"{fmt(multi.get('embedding_risk_roc_auc_mean'))} ± {fmt(multi.get('embedding_risk_roc_auc_std'))}"],
        ["multi-seed coverage-risk AUC mean ± std", f"{fmt(multi.get('coverage_risk_auc_mean'))} ± {fmt(multi.get('coverage_risk_auc_std'))}"],
    ], widths=[3300, 6060])
    add_image(doc, "outputs/robot_3d_demo/robot_3d_embedding_pca.png", "图 2. 合成 3D perception embedding 的 PCA 分布。")
    add_image(doc, "outputs/robot_3d_demo/robot_3d_coverage_risk.png", "图 3. 合成 3D selective perception coverage-risk 曲线。")
    add_para(doc, "规律: 合成 3D 实验的 ROC-AUC 约为 0.80，说明简单几何 embedding 可以提供有用但不完美的可靠性信号。多 seed 结果给出了均值和标准差，比单次实验更适合对导师说明结果不是完全 cherry-pick。")
    add_csv_preview(doc, "outputs/robot_3d_multiseed/robot_3d_multiseed_results.csv", "Synthetic 3D multi-seed 逐 seed 结果")

    doc.add_heading("6. 真实深度图准备与 Smoke Pipeline", level=1)
    add_para(doc, "真实深度图 workflow 的作用，是把项目从纯合成数据推进到可复现的数据处理 pipeline。smoke 数据包含 corridor、obstacle、stairs、tabletop 等场景，用于验证准备、加载、embedding、腐蚀 benchmark 和图像输出是否连通。")
    add_csv_preview(doc, "outputs/real_depth_analysis_smoke/real_depth_scene_summary.csv", "Smoke real-depth scene summary")
    smoke = read_json("outputs/real_depth_corruption_smoke/real_depth_corruption_metrics.json")
    pipe = read_json("outputs/robot_3d_pipeline_smoke/corruption_benchmark/real_depth_corruption_metrics.json")
    add_table(doc, ["实验", "source files", "samples", "ROC-AUC", "coverage-risk AUC"], [
        ["real_depth_corruption_smoke", smoke.get("source_files"), smoke.get("samples"), fmt(smoke.get("embedding_risk_roc_auc_for_corruption_detection")), fmt(smoke.get("coverage_risk_auc"))],
        ["prepared_depth pipeline smoke", pipe.get("source_files"), pipe.get("samples"), fmt(pipe.get("embedding_risk_roc_auc_for_corruption_detection")), fmt(pipe.get("coverage_risk_auc"))],
    ], widths=[3100, 1300, 1200, 1600, 2160])
    add_image(doc, "outputs/real_depth_corruption_smoke/real_depth_corruption_pca.png", "图 4. Smoke real-depth corruption embedding PCA。")
    add_image(doc, "outputs/real_depth_corruption_smoke/real_depth_corruption_coverage_risk.png", "图 5. Smoke real-depth corruption coverage-risk 曲线。")

    doc.add_heading("7. TUM RGB-D 公开数据实验", level=1)
    tum_profile = read_json("outputs/tum_rgbd_freiburg1_desk_robot_3d/profile/real_depth_metrics.json")
    tum_corr = read_json("outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_metrics.json")
    add_para(doc, "TUM RGB-D 是公开 RGB-D 数据集。本项目使用 freiburg1_desk 序列的 300 个深度文件作为公开数据 evidence layer。实验首先分析真实深度图几何 embedding，再对每个深度图生成 clean、gaussian_noise、dropout、quantization、occlusion、tilt_shift 等受控版本。")
    add_table(doc, ["项目", "数值"], [
        ["depth_dir", tum_profile.get("depth_dir")],
        ["profile samples", tum_profile.get("samples")],
        ["scenes", tum_profile.get("scenes")],
        ["corruption source files", tum_corr.get("source_files")],
        ["corruption total samples", tum_corr.get("samples")],
        ["clean samples", tum_corr.get("clean_samples")],
        ["corrupted samples", tum_corr.get("corrupted_samples")],
        ["source-paired ROC-AUC", fmt(tum_corr.get("embedding_risk_roc_auc_for_corruption_detection"))],
        ["scene-conditioned ROC-AUC", fmt(tum_corr.get("scene_conditioned_roc_auc_for_corruption_detection"))],
        ["coverage-risk AUC", fmt(tum_corr.get("coverage_risk_auc"))],
    ], widths=[3300, 6060])
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_robot_3d/profile/real_depth_embedding_pca.png", "图 6. TUM RGB-D depth geometry embedding PCA。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_pca.png", "图 7. TUM RGB-D corruption embedding PCA。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_coverage_risk.png", "图 8. TUM RGB-D source-paired coverage-risk 曲线。")
    add_para(doc, "规律: source-paired ROC-AUC = 1.000，证明如果把每个腐蚀样本和它自己的 clean source 比较，受控腐蚀非常容易检测；但 scene-conditioned ROC-AUC = 0.483，接近随机，说明只拿全局场景参考比较会受到相机运动和深度分布变化干扰。这个负结果很重要，因为它推动后续的 temporal 和 pose-aware 分析。")

    doc.add_heading("8. 时序可靠性、姿态分析和 PCA 描述子", level=1)
    temporal = read_json("outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_metrics.json")
    pose = read_json("outputs/tum_rgbd_freiburg1_desk_pose_sorted/tum_pose_embedding_metrics.json")
    pca = read_json("outputs/tum_rgbd_freiburg1_desk_pca_descriptor/tum_pca_depth_descriptor_metrics.json")
    add_para(doc, "时序可靠性实验尝试解决上一节暴露的问题: 在一个移动相机序列里，正常帧之间本来就会变化，因此风险分数不能只看离全局 clean reference 多远，而应当比较某一帧的腐蚀 shift 是否超过附近 clean 帧的自然变化。")
    add_table(doc, ["指标", "数值"], [
        ["source files", temporal.get("source_files")],
        ["samples", temporal.get("samples")],
        ["temporal window", f"+/- {temporal.get('window')} frames"],
        ["temporal excess ROC-AUC", fmt(temporal.get("temporal_excess_roc_auc_for_corruption_detection"))],
        ["temporal local-distance ROC-AUC", fmt(temporal.get("temporal_local_roc_auc_for_corruption_detection"))],
        ["source-paired ROC-AUC", fmt(temporal.get("source_paired_roc_auc_for_corruption_detection"))],
    ], widths=[3600, 5760])
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_embedding_pca.png", "图 9. TUM temporal depth reliability embedding PCA。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_coverage_risk.png", "图 10. TUM temporal local-reference coverage-risk 曲线。")
    add_para(doc, "姿态分析进一步检查 embedding shift 是否和真实相机位姿变化相关。结果显示全局手工几何特征对旋转几乎不敏感，而保留局部布局的 grid descriptor 对旋转更敏感。")
    add_table(doc, ["描述子", "translation Spearman", "rotation Spearman", "mean embedding shift"], [
        ["global descriptor", fmt(pose.get("global_spearman_translation_vs_embedding_shift")), fmt(pose.get("global_spearman_rotation_vs_embedding_shift")), fmt(pose.get("global_mean_embedding_shift"))],
        ["grid descriptor", fmt(pose.get("grid_spearman_translation_vs_embedding_shift")), fmt(pose.get("grid_spearman_rotation_vs_embedding_shift")), fmt(pose.get("grid_mean_embedding_shift"))],
        ["PCA depth descriptor", fmt(pca.get("spearman_translation_vs_embedding_shift")), fmt(pca.get("spearman_rotation_vs_embedding_shift")), fmt(pca.get("mean_embedding_shift"))],
    ], widths=[2700, 2200, 2200, 2260])
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_pose_sorted/pose_rotation_vs_global_embedding_shift.png", "图 11. TUM global descriptor: rotation vs embedding shift。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_pose_sorted/pose_rotation_vs_grid_embedding_shift.png", "图 12. TUM grid descriptor: rotation vs embedding shift。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_pca_descriptor/pose_rotation_vs_pca_depth_shift.png", "图 13. PCA depth descriptor: rotation vs embedding shift。")
    add_image(doc, "outputs/tum_rgbd_freiburg1_desk_pca_descriptor/pose_translation_vs_pca_depth_shift.png", "图 14. PCA depth descriptor: translation vs embedding shift。")
    add_para(doc, "规律: PCA depth descriptor 的 rotation Spearman = 0.540，高于 grid descriptor 的 0.275 和 global descriptor 的 0.061。这说明轻量 data-driven depth representation 比纯手工全局统计更有潜力捕捉相机运动，适合作为下一阶段 pretrained RGB-D/depth encoder 的前置证据。")

    doc.add_heading("9. 校准、Coverage-Risk 与运行时监控", level=1)
    cal = read_json("outputs/tum_rgbd_calibration/calibration_metrics.json")
    mon = read_json("outputs/tum_rgbd_runtime_monitor/runtime_monitor_metrics.json")
    add_para(doc, "校准分析的目的，是区分“排序能力”和“概率解释能力”。一个风险分数可以很好地区分异常样本，但它的数值本身未必能直接当成真实概率。")
    add_table(doc, ["校准指标", "数值"], [
        ["samples", cal.get("samples")],
        ["positive label count", cal.get("positive_label_count")],
        ["positive label rate", fmt(cal.get("positive_label_rate"))],
        ["ROC-AUC", fmt(cal.get("roc_auc"))],
        ["Average precision", fmt(cal.get("average_precision"))],
        ["ECE-style gap", fmt(cal.get("ece"))],
        ["risk at full coverage", fmt(cal.get("risk_at_full_coverage"))],
        ["risk at 80% coverage", fmt(cal.get("risk_at_80pct_coverage"))],
        ["risk at 50% coverage", fmt(cal.get("risk_at_50pct_coverage"))],
    ], widths=[3600, 5760])
    add_image(doc, "outputs/tum_rgbd_calibration/risk_calibration.png", "图 15. TUM temporal risk score calibration 图。")
    add_image(doc, "outputs/tum_rgbd_calibration/coverage_risk.png", "图 16. TUM temporal coverage-risk 图。")
    add_para(doc, "运行时监控器把连续风险分数转为可审计的系统状态，这一步是把机器学习诊断转向机器人 runtime assurance 的关键桥梁。")
    add_table(doc, ["状态/检查", "数值"], [
        ["samples", mon.get("samples")],
        ["suspect threshold", fmt(mon.get("thresholds", {}).get("suspect"), 6)],
        ["recover threshold", fmt(mon.get("thresholds", {}).get("recover"), 6)],
        ["review threshold", fmt(mon.get("thresholds", {}).get("review"), 6)],
        ["NORMAL", mon.get("state_counts", {}).get("NORMAL")],
        ["SUSPECT", mon.get("state_counts", {}).get("SUSPECT")],
        ["RECOVER", mon.get("state_counts", {}).get("RECOVER")],
        ["HUMAN_REVIEW", mon.get("state_counts", {}).get("HUMAN_REVIEW")],
        ["property violations", mon.get("property_checks", {}).get("violations")],
    ], widths=[3600, 5760])
    add_image(doc, "outputs/tum_rgbd_runtime_monitor/runtime_monitor_trace.png", "图 17. TUM runtime monitor risk trace。")
    add_image(doc, "outputs/tum_rgbd_runtime_monitor/runtime_monitor_state_counts.png", "图 18. TUM runtime monitor state counts。")

    doc.add_heading("10. 轨迹残差与执行可靠性实验", level=1)
    traj = read_json("outputs/trajectory_residual_demo/trajectory_residual_metrics.json")
    tmon = read_json("outputs/trajectory_runtime_monitor/runtime_monitor_metrics.json")
    add_para(doc, "轨迹残差实验把项目从 perception reliability 扩展到 action-outcome reliability。实验生成计划 3D 轨迹和实际观测轨迹，并模拟 normal、drift、oscillation、jump、stuck 五种情况。")
    add_table(doc, ["指标", "数值"], [
        ["samples", traj.get("samples")],
        ["n per type", traj.get("n_per_type")],
        ["trajectory length", traj.get("length")],
        ["failure types", traj.get("failure_types")],
        ["ROC-AUC", fmt(traj.get("roc_auc"))],
        ["average precision", fmt(traj.get("average_precision"))],
    ], widths=[3300, 6060])
    add_image(doc, "outputs/trajectory_residual_demo/trajectory_examples.png", "图 19. planned vs observed trajectory examples。")
    add_image(doc, "outputs/trajectory_residual_demo/trajectory_risk_by_type.png", "图 20. trajectory residual risk by failure type。")
    add_para(doc, "将同一个 runtime monitor 应用于轨迹残差分数后，系统得到 300 个 NORMAL、60 个 SUSPECT、40 个 RECOVER、0 个 HUMAN_REVIEW，并且 no-continue-autonomy-under-recover 检查没有违反。")
    add_image(doc, "outputs/trajectory_runtime_monitor/runtime_monitor_trace.png", "图 21. trajectory runtime monitor risk trace。")
    add_image(doc, "outputs/trajectory_runtime_monitor/runtime_monitor_state_counts.png", "图 22. trajectory runtime monitor state counts。")
    add_csv_preview(doc, "outputs/trajectory_residual_demo/trajectory_residual_summary.csv", "Trajectory residual summary by failure type and severity", max_rows=16)

    doc.add_heading("11. 实验规律、博士申请价值和下一步计划", level=1)
    add_para(doc, "综合所有实验，目前最重要的规律有四条。")
    add_numbered(doc, [
        "受控腐蚀在 source-paired 设置下容易检测，但这不是完整机器人安全证明。",
        "全局 clean reference 会被正常相机运动干扰，TUM scene-conditioned baseline 接近随机，这是一个有研究价值的负结果。",
        "局部时间归一化和局部/学习型描述子能改善可靠性判断，PCA depth descriptor 对旋转的相关性明显更强。",
        "风险分数可以转成 runtime monitor 状态，从而把模型诊断连接到恢复、重新规划和人类审查等机器人安全动作。",
    ])
    add_para(doc, "博士申请价值: 该项目可以被表述为一个 evidence-first research prototype。它展示了从模型训练、embedding 诊断、公开数据实验、风险校准、运行时状态机到执行残差分析的完整研究链条。相比单纯写“我会 PyTorch/LSTM”，它更能体现研究问题意识、实验设计能力和对局限性的清醒认识。")
    add_para(doc, "建议下一步实验: 用 pretrained 或 task-supervised RGB-D/depth encoder 替代 PCA baseline，并把可靠性标签从受控腐蚀扩展到 dataset-native labels，例如 SLAM tracking quality、pose error、navigation progress、tool trajectory error 或 simulator rollout failure。")

    doc.add_heading("12. 附录 A: 关键 CSV 预览", level=1)
    for rel, title in [
        ("outputs/tum_rgbd_freiburg1_desk_temporal/temporal_depth_reliability_summary.csv", "TUM temporal depth reliability summary"),
        ("outputs/tum_rgbd_freiburg1_desk_robot_3d/corruption_benchmark/real_depth_corruption_summary.csv", "TUM real-depth corruption summary"),
        ("outputs/tum_rgbd_calibration/calibration_bins.csv", "TUM calibration bins"),
        ("outputs/tum_rgbd_calibration/coverage_risk.csv", "TUM coverage-risk table"),
        ("outputs/tum_rgbd_freiburg1_desk_pose_sorted/tum_pose_embedding_shifts.csv", "TUM pose embedding shifts preview"),
        ("outputs/tum_rgbd_freiburg1_desk_pca_descriptor/tum_pca_depth_descriptor_shifts.csv", "TUM PCA depth descriptor shifts preview"),
        ("outputs/tum_rgbd_runtime_monitor/runtime_monitor_decisions.csv", "TUM runtime monitor decisions preview"),
        ("outputs/trajectory_runtime_monitor/runtime_monitor_decisions.csv", "Trajectory runtime monitor decisions preview"),
    ]:
        add_csv_preview(doc, rel, title, max_rows=12)

    doc.add_heading("附录 B: 全部图片索引", level=1)
    image_rows = []
    for img in all_files("outputs/**/*.png") + all_files("modules/*.png"):
        image_rows.append([str(img.relative_to(ROOT)), f"{img.stat().st_size / 1024:.1f} KB"])
    add_table(doc, ["图片文件", "大小"], image_rows, widths=[7600, 1760])

    doc.add_heading("附录 C: 指标 JSON 文件索引", level=1)
    json_rows = []
    for p in all_files("outputs/**/*.json"):
        data = read_json(str(p.relative_to(ROOT)))
        json_rows.append([str(p.relative_to(ROOT)), data.get("task", ""), data.get("samples", data.get("source_files", ""))])
    add_table(doc, ["JSON 文件", "task", "samples/source_files"], json_rows, widths=[5600, 2500, 1260])

    doc.add_heading("附录 D: 完整 CSV 数据文件索引", level=1)
    csv_rows = []
    for p in all_files("outputs/**/*.csv") + all_files("modules/*.csv"):
        try:
            df = pd.read_csv(p, nrows=1)
            row_count = sum(1 for _ in open(p, "r", encoding="utf-8", errors="ignore")) - 1
            csv_rows.append([str(p.relative_to(ROOT)), row_count, len(df.columns), ", ".join(map(str, df.columns[:8]))])
        except Exception as exc:
            csv_rows.append([str(p.relative_to(ROOT)), "n/a", "n/a", f"read error: {exc}"])
    add_table(doc, ["CSV 文件", "行数", "列数", "前 8 个列名"], csv_rows, widths=[4500, 900, 900, 3060])

    doc.add_heading("附录 E: 子报告原文整合", level=1)
    for p in all_files("docs/*.md") + all_files("outputs/**/*.md"):
        rel = str(p.relative_to(ROOT))
        add_markdown_report_appendix(doc, rel)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(out)
