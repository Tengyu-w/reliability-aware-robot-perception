"""Add VPPV visual evidence and summary section to the experiment compendium."""

from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "vppv_perception_monitor"
REPORTS_DIR = ROOT / "reports"
FIG_DIR = OUT_DIR / "visual_gallery"
DOCX_NAME = "reliability_aware_robot_perception_complete_experiment_compendium_20260623.docx"


def find_source_docx() -> Path:
    matches = list(Path.home().joinpath("Desktop").rglob(DOCX_NAME))
    if not matches:
        raise FileNotFoundError(f"could not find {DOCX_NAME} under Desktop")
    preferred = [p for p in matches if p.parent.name == "项目合集"]
    return preferred[0] if preferred else matches[0]


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def load_data():
    metrics = load_json(OUT_DIR / "risk_distillation_metrics.json")
    outcome = load_json(OUT_DIR / "risk_outcome_correlation.json")
    feature = pd.read_csv(OUT_DIR / "feature_importance.csv")
    ablation = pd.read_csv(OUT_DIR / "signal_group_ablation.csv")
    state_counts = pd.read_csv(OUT_DIR / "risk_state_counts.csv")
    route_policy = pd.read_csv(OUT_DIR / "vppv_route_policy.csv")
    top_cases = pd.read_csv(OUT_DIR / "top_risk_cases.csv")
    trace = pd.read_csv(OUT_DIR / "risk_trace.csv")
    return metrics, outcome, feature, ablation, state_counts, route_policy, top_cases, trace


def save_architecture_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(12, 6.2))
    ax.axis("off")
    boxes = [
        ("VPPV visual front end\nsegmentation / depth / perceptual regressor / physical state", 0.05, 0.68, 0.28, 0.20, "#d9ecff"),
        ("Reliability signals\ndepth corruption\ntemporal excess\nembedding shift\ntrajectory residual\ncalibration gap", 0.38, 0.62, 0.24, 0.32, "#e8f5e9"),
        ("Distilled model\nLogistic Regression\nRandom Forest\nDecision Tree\noutput: visual_state_risk", 0.68, 0.62, 0.24, 0.32, "#fff3cd"),
        ("Runtime routes\nNORMAL: continue\nSUSPECT: re-perceive\nRECOVER: replan / backup\nHUMAN_REVIEW: surgeon review", 0.38, 0.14, 0.54, 0.28, "#f8d7da"),
    ]
    for text, x, y, w, h, color in boxes:
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="#243447", linewidth=1.5))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=12, linespacing=1.25)

    arrows = [
        ((0.33, 0.78), (0.38, 0.78)),
        ((0.62, 0.78), (0.68, 0.78)),
        ((0.80, 0.62), (0.66, 0.42)),
        ((0.50, 0.62), (0.50, 0.42)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=2, color="#243447"))

    ax.text(
        0.5,
        0.04,
        "Core idea: a downstream VPPV policy should not trust visual state blindly; unreliable state triggers re-perception, recovery, or human review.",
        ha="center",
        va="center",
        fontsize=11,
        color="#333333",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_dashboard(feature: pd.DataFrame, ablation: pd.DataFrame, state_counts: pd.DataFrame, outcome: dict, path: Path):
    fig, axes = plt.subplots(2, 2, figsize=(13, 8))
    axes = axes.ravel()

    f = feature.sort_values("importance", ascending=True)
    axes[0].barh(f["feature"], f["importance"], color="#4c78a8")
    axes[0].set_title("Feature attribution")
    axes[0].set_xlabel("importance")

    a = ablation.sort_values("teacher_roc_auc", ascending=True)
    axes[1].barh(a["signal_group"], a["teacher_roc_auc"], color="#59a14f")
    axes[1].set_title("Signal-group ablation")
    axes[1].set_xlabel("teacher ROC-AUC")
    axes[1].set_xlim(0, 1.05)

    colors = ["#2ca02c", "#ff7f0e", "#d62728", "#9467bd"]
    axes[2].bar(state_counts["monitor_state"], state_counts["count"], color=colors)
    axes[2].set_title("Runtime state counts")
    axes[2].set_ylabel("count")
    axes[2].tick_params(axis="x", rotation=20)

    corr_labels = [
        "trajectory\nresidual",
        "temporal\nexcess",
        "embedding\nshift",
        "progress\nstagnation",
        "coverage\nrisk",
    ]
    corr_values = [
        outcome.get("spearman_risk_vs_trajectory_residual", 0.0),
        outcome.get("spearman_risk_vs_temporal_excess_score", 0.0),
        outcome.get("spearman_risk_vs_embedding_shift", 0.0),
        outcome.get("spearman_risk_vs_progress_stagnation_score", 0.0),
        outcome.get("spearman_risk_vs_coverage_risk_score", 0.0),
    ]
    axes[3].bar(corr_labels, corr_values, color="#b07aa1")
    axes[3].axhline(0, color="#333333", linewidth=0.8)
    axes[3].set_title("Outcome-linked correlations")
    axes[3].set_ylabel("Spearman rho")
    axes[3].set_ylim(-0.2, 0.7)

    fig.suptitle("VPPV Perception-State Monitor: What Is Visible At A Glance", fontsize=16)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_route_policy_flow(state_counts: pd.DataFrame, route_policy: pd.DataFrame, path: Path):
    fig, ax = plt.subplots(figsize=(12, 5.5))
    ax.axis("off")
    colors = {
        "NORMAL": "#d9ead3",
        "SUSPECT": "#fff2cc",
        "RECOVER": "#f4cccc",
        "HUMAN_REVIEW": "#eadcf8",
    }
    x_positions = [0.04, 0.28, 0.52, 0.76]
    for idx, row in route_policy.iterrows():
        state = row["monitor_state"]
        count_row = state_counts[state_counts["monitor_state"] == state]
        count = int(count_row["count"].iloc[0]) if not count_row.empty else 0
        fraction = float(count_row["fraction"].iloc[0]) if not count_row.empty else 0.0
        x = x_positions[idx]
        ax.add_patch(plt.Rectangle((x, 0.35), 0.20, 0.34, facecolor=colors[state], edgecolor="#333333", linewidth=1.4))
        ax.text(x + 0.10, 0.60, state, ha="center", va="center", fontsize=13, fontweight="bold")
        ax.text(x + 0.10, 0.51, f"n={count} ({fraction:.1%})", ha="center", va="center", fontsize=11)
        ax.text(x + 0.10, 0.42, row["vppv_action"], ha="center", va="center", fontsize=10, wrap=True)
        if idx < 3:
            ax.annotate("", xy=(x + 0.235, 0.52), xytext=(x + 0.205, 0.52), arrowprops=dict(arrowstyle="->", lw=1.7))

    ax.text(
        0.5,
        0.15,
        "Route policy interpretation: high-risk visual state is not just flagged; it is mapped to a concrete VPPV action.",
        ha="center",
        fontsize=12,
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_top_cases_heatmap(top_cases: pd.DataFrame, path: Path, top_k: int = 18):
    driver_cols = ["depth", "temporal", "embedding", "trajectory", "progress", "calibration", "coverage"]
    rows = top_cases.head(top_k).copy()
    matrix = np.zeros((len(rows), len(driver_cols)))
    for i, row in rows.iterrows():
        for rank in [1, 2, 3]:
            name = row[f"top_driver_{rank}"]
            score = row[f"top_driver_{rank}_score"]
            if name in driver_cols:
                matrix[list(rows.index).index(i), driver_cols.index(name)] = score

    fig, ax = plt.subplots(figsize=(11, 6.2))
    im = ax.imshow(matrix, aspect="auto", cmap="YlOrRd", vmin=0, vmax=1)
    ax.set_xticks(np.arange(len(driver_cols)))
    ax.set_xticklabels(driver_cols, rotation=30, ha="right")
    labels = [
        f"{row.monitor_state} | {row.corruption} | {row.trajectory_failure_type}"
        for row in rows.itertuples()
    ]
    ax.set_yticks(np.arange(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_title("Top High-Risk Cases: Driver Heatmap")
    cbar = fig.colorbar(im, ax=ax)
    cbar.set_label("normalized driver score")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_state_strip(trace: pd.DataFrame, path: Path):
    mapping = {"NORMAL": 0, "SUSPECT": 1, "RECOVER": 2, "HUMAN_REVIEW": 3}
    colors = ["#2ca02c", "#ff7f0e", "#d62728", "#9467bd"]
    arr = trace["risk_monitor_state"].map(mapping).to_numpy()[None, :]
    fig, ax = plt.subplots(figsize=(12, 1.9))
    cmap = plt.matplotlib.colors.ListedColormap(colors)
    ax.imshow(arr, aspect="auto", cmap=cmap, vmin=0, vmax=3)
    ax.set_yticks([])
    ax.set_xlabel("sample index")
    ax.set_title("Runtime State Strip Across 1,800 Samples")
    handles = [
        plt.Line2D([0], [0], marker="s", color="w", label=state, markerfacecolor=colors[idx], markersize=10)
        for state, idx in mapping.items()
    ]
    ax.legend(handles=handles, ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.35), frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def generate_visuals():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    metrics, outcome, feature, ablation, state_counts, route_policy, top_cases, trace = load_data()
    paths = {
        "architecture": FIG_DIR / "vppv_monitor_architecture.png",
        "dashboard": FIG_DIR / "vppv_visual_dashboard.png",
        "route_flow": FIG_DIR / "vppv_route_policy_flow.png",
        "top_cases": FIG_DIR / "vppv_top_risk_case_heatmap.png",
        "state_strip": FIG_DIR / "vppv_state_strip.png",
    }
    save_architecture_diagram(paths["architecture"])
    save_dashboard(feature, ablation, state_counts, outcome, paths["dashboard"])
    save_route_policy_flow(state_counts, route_policy, paths["route_flow"])
    save_top_cases_heatmap(top_cases, paths["top_cases"])
    save_state_strip(trace, paths["state_strip"])
    return paths, metrics, outcome, feature, ablation, state_counts, route_policy, top_cases


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill="E8EEF5"):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_figure(doc: Document, image_path: Path, caption: str, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_heading(doc: Document, text: str, level: int):
    return doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_metric_table(doc: Document, metrics: dict, outcome: dict):
    rows = [
        ("Selected distilled model", metrics["selected_model"]),
        ("Teacher ROC-AUC, selected model", f"{metrics['model_metrics'][metrics['selected_model']]['teacher_roc_auc']:.3f}"),
        ("Teacher average precision", f"{metrics['model_metrics'][metrics['selected_model']]['teacher_average_precision']:.3f}"),
        ("State counts", "1350 NORMAL / 433 SUSPECT / 17 RECOVER / 0 HUMAN_REVIEW"),
        ("Risk vs trajectory residual", f"Spearman rho {outcome['spearman_risk_vs_trajectory_residual']:.3f}"),
        ("Risk vs temporal excess", f"Spearman rho {outcome['spearman_risk_vs_temporal_excess_score']:.3f}"),
        ("Top 10% risk capture", f"{outcome['top10_risk_recover_or_review_capture']:.3f} of RECOVER/HUMAN_REVIEW states"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Metric", bold=True)
    set_cell_text(table.rows[0].cells[1], "Value", bold=True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key)
        set_cell_text(cells[1], value)
    style_table(table)


def add_route_table(doc: Document, route_policy: pd.DataFrame):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Monitor state", "中文", "VPPV action", "中文动作"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
    for row in route_policy.itertuples():
        cells = table.add_row().cells
        values = [row.monitor_state, row.state_zh, row.vppv_action, row.action_zh]
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table)


def add_top_case_table(doc: Document, top_cases: pd.DataFrame):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["State", "Corruption", "Failure type", "Risk", "Explanation"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
    for row in top_cases.head(5).itertuples():
        cells = table.add_row().cells
        values = [
            row.monitor_state,
            row.corruption,
            row.trajectory_failure_type,
            f"{row.visual_state_risk:.3f}",
            row.explanation,
        ]
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    style_table(table)


def append_vppv_section(docx_path: Path, output_path: Path):
    metrics, outcome, feature, ablation, state_counts, route_policy, top_cases, _ = load_data()
    visuals = {
        "architecture": FIG_DIR / "vppv_monitor_architecture.png",
        "dashboard": FIG_DIR / "vppv_visual_dashboard.png",
        "route_flow": FIG_DIR / "vppv_route_policy_flow.png",
        "top_cases": FIG_DIR / "vppv_top_risk_case_heatmap.png",
        "state_strip": FIG_DIR / "vppv_state_strip.png",
    }
    missing = [path for path in visuals.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "visual gallery is missing; run tools/generate_vppv_visual_gallery.py first: "
            + ", ".join(str(path) for path in missing)
        )
    doc = Document(str(docx_path))

    doc.add_page_break()
    add_heading(doc, "17. VPPV-Style Surgical Autonomy Upgrade", 1)
    p = doc.add_paragraph()
    p.add_run("New contribution. ").bold = True
    p.add_run(
        "The reliability-aware robot perception project is upgraded into a visual-front-end monitor for VPPV-style surgical autonomy. "
        "VPPV depends on segmentation masks, depth maps, perceptual regressors, and physical state vectors; this section shows how the project can monitor whether those visual-state inputs are reliable enough for a downstream policy."
    )

    add_heading(doc, "17.1 What Can Be Seen Directly", 2)
    add_bullet(doc, "Depth and embedding anomalies can be visualized as temporal risk traces and state strips.")
    add_bullet(doc, "Feature attribution shows which signals drive visual_state_risk, rather than reporting a black-box scalar.")
    add_bullet(doc, "Signal-group ablation shows which evidence family is useful: depth, temporal, embedding, trajectory, calibration, or coverage.")
    add_bullet(doc, "Route policy diagrams show what the surgical autonomy system should do after a risk state is detected.")
    add_bullet(doc, "Top-risk case heatmaps make individual high-risk examples inspectable.")

    add_figure(
        doc,
        visuals["architecture"],
        "Figure 17.1. VPPV-style reliability monitor architecture: visual front end -> reliability signals -> distilled risk -> runtime route.",
    )

    add_heading(doc, "17.2 Visual Dashboard", 2)
    add_metric_table(doc, metrics, outcome)
    add_figure(
        doc,
        visuals["dashboard"],
        "Figure 17.2. One-page dashboard for feature attribution, signal ablation, runtime states, and outcome-linked correlations.",
    )

    add_heading(doc, "17.3 Feature Attribution And High-Risk Explanations", 2)
    p = doc.add_paragraph()
    p.add_run("Interpretation. ").bold = True
    p.add_run(
        "The selected Random Forest is not presented only by ROC-AUC. It is also inspected by feature importance and top-risk case explanations. "
        "In this run, embedding_shift, trajectory_residual, and progress_slope are the strongest student-model contributors."
    )
    add_figure(
        doc,
        OUT_DIR / "feature_importance.png",
        "Figure 17.3. Feature-level evidence for the distilled visual_state_risk model.",
    )
    add_top_case_table(doc, top_cases)
    add_figure(
        doc,
        visuals["top_cases"],
        "Figure 17.4. Top high-risk cases shown as a driver heatmap. Rows combine monitor state, visual corruption, and trajectory failure type.",
    )

    add_heading(doc, "17.4 Signal-Group Ablation", 2)
    p = doc.add_paragraph()
    p.add_run("Question answered. ").bold = True
    p.add_run(
        "If VPPV front-end state is unreliable, which signal family helps most? The ablation compares depth-only, temporal-only, embedding-only, trajectory-only, calibration-only, all student features, and heavier all-signal variants."
    )
    add_figure(
        doc,
        OUT_DIR / "signal_group_ablation.png",
        "Figure 17.5. Signal-group ablation for visual-state risk distillation.",
    )

    add_heading(doc, "17.5 Runtime Route Evaluation", 2)
    add_route_table(doc, route_policy)
    add_figure(
        doc,
        visuals["route_flow"],
        "Figure 17.6. VPPV action route policy. Each monitor state maps to a concrete autonomy action.",
    )
    add_figure(
        doc,
        visuals["state_strip"],
        "Figure 17.7. Runtime state strip across 1,800 samples. Green is NORMAL, orange is SUSPECT, red is RECOVER, purple is HUMAN_REVIEW.",
    )
    add_figure(
        doc,
        OUT_DIR / "risk_trace.png",
        "Figure 17.8. Full visual_state_risk trace with thresholds and component signals.",
    )

    add_heading(doc, "17.6 Evidence, Limitations, And Next Step", 2)
    add_bullet(doc, "Confirmed: the code now produces risk distillation metrics, state counts, attribution plots, ablation plots, route policy, and outcome-linked validation.")
    add_bullet(doc, "Suggested: the same monitor can serve as a VPPV front-end reliability score, re-perception trigger, recovery trigger, or human-review trigger.")
    add_bullet(doc, "Limitation: current data are controlled RGB-D corruptions and aligned trajectory-residual proxies, not paired surgical VPPV rollouts.")
    add_bullet(doc, "Next experiment: replace proxy residuals with surgical-tool tracking, simulator rollouts, segmentation-mask quality, or VPPV policy failure labels.")

    add_heading(doc, "17.7 Reproduction Command", 2)
    p = doc.add_paragraph()
    r = p.add_run("python modules/run_vppv_perception_monitor.py")
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.save(str(output_path))
    return output_path


def main():
    source = find_source_docx()
    output = source.with_name(source.stem + "_vppv_visual_upgrade_20260625.docx")
    append_vppv_section(source, output)
    print(output)


if __name__ == "__main__":
    main()
